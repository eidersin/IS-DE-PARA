# Fica em: src/support_generator.py

from docx import Document
from docx.shared import Inches
from typing import Dict, Any, List
import logging

# ==============================================================================
# FUNÇÕES AUXILIARES DE FORMATAÇÃO
# ==============================================================================

def add_info(doc: Document, label: str, value: Any):
    """
    Adiciona um parágrafo formatado com um rótulo em negrito e um valor.
    Ignora valores nulos ou vazios para manter o documento limpo.
    """
    if value is not None and value != '' and value != []:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        # Se o valor for uma lista, junta os itens com vírgula
        if isinstance(value, list):
            p.add_run(', '.join(map(str, value)))
        else:
            p.add_run(str(value))

def add_bool_info(doc: Document, label: str, condition: bool):
    """Adiciona uma informação de Sim/Não baseada em um valor booleano."""
    if condition is not None:
        add_info(doc, label, "Sim" if condition else "Não")

def add_list_as_bullets(doc: Document, label: str, items: List[str]):
    """Adiciona uma lista de itens como marcadores (bullets) no documento."""
    if items:
        doc.add_paragraph(f'{label}:').bold = True
        for item in items:
            doc.add_paragraph(str(item), style='List Bullet')

def add_table_from_list(doc: Document, data_list: List[Dict], headers: Dict[str, str]):
    """
    Cria uma tabela a partir de uma lista de dicionários.

    Args:
        doc (Document): O objeto do documento docx.
        data_list (List[Dict]): A lista de dados (cada dict é uma linha).
        headers (Dict[str, str]): Um dicionário mapeando chave do JSON para o cabeçalho da tabela.
    """
    if not data_list:
        return

    header_keys = list(headers.keys())
    header_values = list(headers.values())

    table = doc.add_table(rows=1, cols=len(header_values))
    table.style = 'Table Grid'
    table.autofit = True

    # Preenche o cabeçalho
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(header_values):
        hdr_cells[i].text = header_text

    # Preenche as linhas de dados
    for item in data_list:
        row_cells = table.add_row().cells
        for i, key in enumerate(header_keys):
            # Usa .get() para evitar erro se a chave não existir
            cell_value = item.get(key, 'N/A')
            row_cells[i].text = str(cell_value)

# ==============================================================================
# FUNÇÃO PRINCIPAL PARA GERAR DOCUMENTO DE FOLHA DE PAGAMENTO
# ==============================================================================

def create_folha_pagamento_support_document(data: Dict[str, Any], output_path: str):
    """
    Cria um documento de suporte .docx para Folha de Pagamento a partir da
    estrutura de dados do 'super_schema_folha.json'.
    """
    if not data:
        logging.warning("Dados para Folha de Pagamento estão vazios. Documento não gerado.")
        return

    try:
        doc = Document()
        doc.add_heading('Diagnóstico de Folha de Pagamento - Análise da CCT', level=0)

        # 1. Identificação do Documento
        if 'identificacao_documento' in data:
            doc.add_heading('1. Identificação do Documento', level=1)
            id_doc = data['identificacao_documento']
            if id_doc.get('sindicato_empregados'):
                doc.add_heading('Sindicato dos Empregados', level=2)
                sind_emp = id_doc['sindicato_empregados']
                add_info(doc, 'Nome', sind_emp.get('nome'))
                add_info(doc, 'CNPJ', sind_emp.get('cnpj'))
                add_info(doc, 'Representante', sind_emp.get('representante'))
            if id_doc.get('sindicato_patronal'):
                doc.add_heading('Sindicato Patronal', level=2)
                sind_pat = id_doc['sindicato_patronal']
                add_info(doc, 'Nome', sind_pat.get('nome'))
                add_info(doc, 'CNPJ', sind_pat.get('cnpj'))
                add_info(doc, 'Representante', sind_pat.get('representante'))
            add_info(doc, 'Vigência do Instrumento', id_doc.get('vigencia'))
            add_info(doc, 'Data Base da Categoria', id_doc.get('data_base_categoria'))
            add_list_as_bullets(doc, 'Abrangência Territorial', id_doc.get('abrangencia_territorial'))

        # 2. Pisos Salariais
        if data.get('pisos_salariais'):
            doc.add_heading('2. Pisos Salariais', level=1)
            add_table_from_list(doc, data['pisos_salariais'], {
                'cargo_funcao': 'Cargo/Função',
                'valor_piso': 'Piso (R$)',
                'observacoes': 'Observações'
            })
        
        # 3. Reajuste Salarial
        if 'reajuste_salarial' in data:
            doc.add_heading('3. Reajuste Salarial', level=1)
            reajuste = data['reajuste_salarial']
            add_info(doc, 'Percentual', reajuste.get('percentual_reajuste'))
            add_info(doc, 'Data de Aplicação', reajuste.get('data_aplicacao'))
            add_info(doc, 'Regras de Proporcionalidade', reajuste.get('regras_proporcionalidade_admitidos'))
            add_info(doc, 'Compensações Permitidas', reajuste.get('compensacoes_permitidas'))

        # Mapeamento para seções genéricas
        section_map = {
            'beneficios': '4. Benefícios Concedidos',
            'adicionais_e_gratificacoes': '5. Adicionais e Gratificações',
            'jornada_de_trabalho': '6. Jornada de Trabalho',
            'estabilidade_provisoria': '7. Regras de Estabilidade',
            'rescisao_contratual': '8. Regras de Rescisão Contratual',
            'saude_e_seguranca': '9. Saúde e Segurança do Trabalhador',
            'contribuicoes_sindicais_empregados': '10. Contribuições Sindicais (Empregados)'
        }

        for key, title in section_map.items():
            if data.get(key):
                doc.add_heading(title, level=1)
                section_data = data[key]
                for sub_key, sub_value in section_data.items():
                    doc.add_heading(sub_key.replace("_", " ").title(), level=2)
                    if isinstance(sub_value, dict):
                        for item_key, item_value in sub_value.items():
                            if isinstance(item_value, dict):
                                doc.add_heading(item_key.replace("_", " ").title(), level=3)
                                for inner_key, inner_value in item_value.items():
                                    add_info(doc, inner_key.replace("_", " ").title(), inner_value)
                            else:
                                add_info(doc, item_key.replace("_", " ").title(), item_value)
                    else:
                        add_info(doc, sub_key.replace("_", " ").title(), sub_value)

        doc.save(output_path)
        logging.info(f"Documento de Folha de Pagamento gerado com sucesso em: {output_path}")

    except Exception as e:
        logging.error(f"Falha ao gerar o documento de Folha de Pagamento: {e}", exc_info=True)

# ==============================================================================
# FUNÇÃO PRINCIPAL PARA GERAR DOCUMENTO DE INTERFACE CONTÁBIL
# ==============================================================================

def create_interface_contabil_support_doc(data: Dict[str, Any], output_path: str):
    """
    Cria um documento de suporte .docx para Interface Contábil a partir da
    estrutura de dados do 'super_schema_contabil.json'.
    """
    if not data:
        logging.warning("Dados para Interface Contábil estão vazios. Documento não gerado.")
        return

    try:
        doc = Document()
        doc.add_heading('Diagnóstico de Interface Contábil', level=0)

        # 1. Informações Gerais
        if 'informacoes_gerais' in data:
            doc.add_heading('1. Informações Gerais', level=1)
            info = data['informacoes_gerais']
            add_info(doc, 'Nome do Cliente', info.get('nome_cliente'))
            add_info(doc, 'Responsável pela Aprovação', info.get('responsavel_aprovacao_diagnostico'))
            add_info(doc, 'Sistema Contábil Utilizado', info.get('sistema_contabil_utilizado'))

        # 2. Especificações do Layout do Arquivo
        if data.get('especificacoes_layout_arquivo'):
            doc.add_heading('2. Especificações do Layout do Arquivo', level=1)
            layout = data['especificacoes_layout_arquivo']
            add_info(doc, 'Formato do Arquivo', layout.get('formato_arquivo'))
            add_info(doc, 'Observações', layout.get('observacoes_formato'))
            if layout.get('campos_do_layout'):
                doc.add_paragraph('Estrutura dos Campos:', style='Heading 2')
                add_table_from_list(doc, layout['campos_do_layout'], {
                    'nome_campo': 'Nome do Campo',
                    'formato': 'Formato',
                    'regra_especifica': 'Regra Específica'
                })
        
        # 3. Conceitos Gerais da Contabilização
        if data.get('conceitos_gerais_contabilizacao'):
            doc.add_heading('3. Conceitos Gerais da Contabilização', level=1)
            conceitos = data['conceitos_gerais_contabilizacao']
            if conceitos.get('integracao_empresa'):
                doc.add_heading('Integração por Empresa', level=2)
                integracao = conceitos['integracao_empresa']
                add_bool_info(doc, 'Códigos de Empresa Iguais na Integração', integracao.get('codigos_empresa_iguais_integracao'))
                add_info(doc, 'Regra para Códigos Diferentes', integracao.get('regra_codigos_diferentes'))
                if integracao.get('lista_empresas'):
                    doc.add_paragraph('Empresas Mapeadas:', style='Heading 3')
                    add_table_from_list(doc, integracao['lista_empresas'], {
                        'codigo_contabil': 'Código Contábil', 'razao_social': 'Razão Social', 'cnpj': 'CNPJ'
                    })
            # Adicione outras seções de 'conceitos' aqui...

        # 4. Configuração da Conta Contábil
        if data.get('configuracao_conta_contabil'):
            doc.add_heading('4. Configuração da Conta Contábil', level=1)
            config = data['configuracao_conta_contabil']
            if config.get('estrutura_conta'):
                doc.add_heading('Estrutura da Conta', level=2)
                estrutura = config['estrutura_conta']
                add_info(doc, 'Tamanho do Código', estrutura.get('tamanho_codigo'))
                add_info(doc, 'Máscara', estrutura.get('mascara'))
                add_info(doc, 'Tamanho da Descrição', estrutura.get('tamanho_descricao_conta'))
            if config.get('parametros_lancamento'):
                doc.add_heading('Parâmetros de Lançamento', level=2)
                params = config['parametros_lancamento']
                add_list_as_bullets(doc, 'Conceitos Base para Lançamento', params.get('conceitos_base_lancamento'))
                add_info(doc, 'Tipo de Lançamento', params.get('tipo_lancamento'))
            # Adicione outras seções de 'configuracao_conta_contabil' aqui...

        # 5. Configuração do Rateio Contábil
        if data.get('configuracao_rateio_contabil'):
            doc.add_heading('5. Configuração do Rateio Contábil', level=1)
            rateio = data['configuracao_rateio_contabil']
            add_bool_info(doc, 'Possui Rateio', rateio.get('possui_rateio'))
            if rateio.get('possui_rateio'):
                add_list_as_bullets(doc, 'Critérios de Apuração de Custo', rateio.get('criterios_apuracao_custo'))
                if rateio.get('lancamento_rateio'):
                    doc.add_heading('Lançamento do Rateio', level=2)
                    lanc_rateio = rateio['lancamento_rateio']
                    add_info(doc, 'Método de Lançamento', lanc_rateio.get('metodo_lancamento'))
                    add_info(doc, 'Regra Específica', lanc_rateio.get('regra_especifica_metodo'))
                if rateio.get('calculo_automatico_rateio'):
                    doc.add_heading('Cálculo Automático', level=2)
                    calc_auto = rateio['calculo_automatico_rateio']
                    add_bool_info(doc, 'Habilitado', calc_auto.get('habilitado'))
                    add_info(doc, 'Regra', calc_auto.get('regra_calculo_automatico'))
        
        doc.save(output_path)
        logging.info(f"Documento de Interface Contábil gerado com sucesso em: {output_path}")

    except Exception as e:
        logging.error(f"Falha ao gerar o documento de Interface Contábil: {e}", exc_info=True)