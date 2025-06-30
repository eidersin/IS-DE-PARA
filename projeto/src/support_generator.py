# support_generator.py (versão atualizada para exibir o Super Schema)

from docx import Document

def add_info(doc, label, value):
    """Adiciona um parágrafo formatado se o valor não for nulo."""
    if value:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))

def add_bool_info(doc, label, condition):
    """Adiciona uma informação de Sim/Não."""
    if condition is not None:
        add_info(doc, label, "Sim" if condition else "Não")

def create_support_document(data: dict, output_path: str):
    if not data:
        print("Dados de entrada vazios.")
        return

    doc = Document()
    doc.add_heading('Documento de Suporte - Análise da Convenção Coletiva', level=0)

    # --- Identificação do Documento ---
    doc.add_heading('1. Identificação do Documento', level=1)
    id_doc = data.get("identificacao_documento", {})
    add_info(doc, 'Vigência', id_doc.get("vigencia"))
    add_info(doc, 'Data Base', id_doc.get("data_base_categoria"))
    sindicatos = id_doc.get("sindicatos_participantes", [])
    if sindicatos:
        doc.add_paragraph("Sindicatos Envolvidos:").bold = True
        for s in sindicatos:
            doc.add_paragraph(f"- {s.get('nome')} (CNPJ: {s.get('cnpj')}, Rep: {s.get('representante')})", style='List Bullet')

    # --- Pisos Salariais ---
    pisos = data.get("pisos_salariais", [])
    if pisos:
        doc.add_heading('2. Pisos Salariais', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cargo'
        hdr_cells[1].text = 'Piso (R$)'
        for piso in pisos:
            row_cells = table.add_row().cells
            row_cells[0].text = piso.get('cargo')
            row_cells[1].text = str(piso.get('valor_piso'))
            
    # --- Reajuste Salarial ---
    reajuste = data.get("reajuste_salarial", {})
    if reajuste:
        doc.add_heading('3. Reajuste Salarial', level=1)
        add_info(doc, 'Percentual', reajuste.get("percentual_reajuste"))
        add_info(doc, 'Detalhes', reajuste.get("detalhes"))

    # --- Benefícios ---
    beneficios = data.get("beneficios", {})
    if beneficios:
        doc.add_heading('4. Benefícios', level=1)
        for nome_beneficio, detalhes in beneficios.items():
            if detalhes and detalhes.get("possui"):
                doc.add_heading(nome_beneficio.replace("_", " ").title(), level=2)
                for chave, valor in detalhes.items():
                     add_info(doc, chave.replace("_", " ").title(), valor)

    # --- Adicionais ---
    adicionais = data.get("adicionais", {})
    if adicionais:
        doc.add_heading('5. Adicionais', level=1)
        for nome_adicional, detalhes in adicionais.items():
            if detalhes:
                doc.add_heading(nome_adicional.replace("_", " ").title(), level=2)
                for chave, valor in detalhes.items():
                     add_info(doc, chave.replace("_", " ").title(), valor)
    
    # --- Jornada de Trabalho ---
    jornada = data.get("jornada_trabalho", {})
    if jornada:
        doc.add_heading('6. Jornada de Trabalho', level=1)
        for tipo_jornada, detalhes in jornada.items():
            if detalhes:
                doc.add_heading(tipo_jornada.replace("_", " ").title(), level=2)
                for chave, valor in detalhes.items():
                    add_bool_info(doc, chave.replace("_", " ").title(), valor) if isinstance(valor, bool) else add_info(doc, chave.replace("_", " ").title(), valor)

    # --- Estabilidade ---
    estabilidade = data.get("estabilidade", {})
    if estabilidade:
        doc.add_heading('7. Regras de Estabilidade', level=1)
        for tipo_estabilidade, detalhes in estabilidade.items():
            if detalhes and detalhes.get("possui"):
                doc.add_heading(tipo_estabilidade.replace("_", " ").title(), level=2)
                for chave, valor in detalhes.items():
                     add_info(doc, chave.replace("_", " ").title(), valor)

    doc.save(output_path)
    print(f"Documento de suporte (versão detalhada) gerado com sucesso em: {output_path}")